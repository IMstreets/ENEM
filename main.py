"""Sistema Gerador de Compilados de Questões do ENEM
Aplicação principal para geração de documentos Word com questões filtradas.
"""

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
import requests
from io import BytesIO
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import logging
from dataclasses import dataclass

from ListaProva import provas_regulares
from config_docx import configurar_documento, adicionar_paragrafo_formatado
from entradas import entradas

# Configuração de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class Questao:
    """Representa uma questão do ENEM com todos seus dados."""
    numero: int
    ano: int
    posicao: int
    nota: float
    habilidade: str
    gabarito: str
    contexto: Optional[str] = None
    referencias: Optional[str] = None
    introducao_alternativas: Optional[str] = None
    alternativas: List[Dict] = None
    arquivos: List[str] = None
    
    def __post_init__(self):
        """Inicializa listas vazias se não fornecidas."""
        if self.alternativas is None:
            self.alternativas = []
        if self.arquivos is None:
            self.arquivos = []


class GeradorDocumento:
    """Classe responsável pela geração do documento Word."""
    
    # Configurações de formatação
    LARGURA_IMAGEM_CONTEXTO = 9  # cm
    LARGURA_IMAGEM_QUESTAO = 7   # cm
    LARGURA_IMAGEM_ALTERNATIVA = 5  # cm
    TAMANHO_FONTE_REFERENCIA = 7  # pt
    
    def __init__(self):
        """Inicializa o gerador de documentos."""
        self.doc = None
        self.imagens_cache = {}
    
    def criar_documento(self, questoes: List[Questao], anos: List[str]) -> Document:
        """
        Cria documento Word com as questões fornecidas.
        
        Args:
            questoes: Lista de questões a incluir
            anos: Lista de anos das provas
            
        Returns:
            Documento Word configurado
        """
        self.doc = Document()
        self.doc = configurar_documento(self.doc)
        
        # Adiciona questões
        for questao in questoes:
            self._adicionar_questao(questao)
        
        # Adiciona gabarito
        self._adicionar_gabarito(questoes)
        
        return self.doc
    
    def _adicionar_questao(self, questao: Questao):
        """Adiciona uma questão completa ao documento."""
        # Título da questão
        self._adicionar_titulo_questao(questao)
        
        # Contexto e imagens
        imagens_processadas = set()
        if questao.contexto:
            imagens_processadas = self._processar_contexto(questao.contexto)
        
        # Imagens adicionais
        self._adicionar_imagens_questao(questao.arquivos, imagens_processadas)
        
        # Referências
        if questao.referencias:
            self._adicionar_referencias(questao.referencias)
        
        # Introdução das alternativas
        if questao.introducao_alternativas:
            adicionar_paragrafo_formatado(self.doc, questao.introducao_alternativas)
        
        # Alternativas
        self._adicionar_alternativas(questao.alternativas)
    
    def _adicionar_titulo_questao(self, questao: Questao):
        """Adiciona título formatado da questão."""
        titulo = (
            f"{questao.numero}. (ENEM {questao.ano} - Questão {questao.posicao}) "
            f"(H{questao.habilidade} - {questao.nota:.2f})"
        )
        
        par = self.doc.add_paragraph()
        run = par.add_run(titulo)
        run.bold = True
    
    def _processar_contexto(self, contexto: str) -> set:
        """
        Processa contexto da questão, separando texto e imagens.
        
        Args:
            contexto: Texto do contexto com possíveis imagens em markdown
            
        Returns:
            Set de URLs de imagens processadas
        """
        # Pattern para identificar imagens em markdown
        pattern = r'!\[[^\]]*\]\(([^)]+)\)'
        partes = re.split(pattern, contexto)
        
        imagens_processadas = set()
        
        for i, parte in enumerate(partes):
            if i % 2 == 0:
                # Parte de texto
                if parte.strip():
                    adicionar_paragrafo_formatado(self.doc, parte.strip())
            else:
                # URL de imagem
                img_url = parte.strip()
                imagens_processadas.add(img_url)
                self._adicionar_imagem(img_url, self.LARGURA_IMAGEM_CONTEXTO)
        
        return imagens_processadas
    
    def _adicionar_imagens_questao(self, urls_imagens: List[str], 
                                  imagens_ja_processadas: Optional[set] = None):
        """Adiciona imagens da questão que ainda não foram processadas."""
        if imagens_ja_processadas is None:
            imagens_ja_processadas = set()
        
        for img_url in urls_imagens:
            if img_url not in imagens_ja_processadas:
                self._adicionar_imagem(img_url, self.LARGURA_IMAGEM_QUESTAO)
    
    def _adicionar_imagem(self, url: str, largura_cm: float) -> bool:
        """
        Adiciona imagem ao documento.
        
        Args:
            url: URL da imagem
            largura_cm: Largura da imagem em centímetros
            
        Returns:
            True se sucesso, False caso contrário
        """
        try:
            # Verifica cache
            if url in self.imagens_cache:
                img_data = self.imagens_cache[url]
                img_data.seek(0)  # Reseta posição do buffer
            else:
                # Baixa imagem
                resp = requests.get(url, timeout=10)
                resp.raise_for_status()
                img_data = BytesIO(resp.content)
                self.imagens_cache[url] = img_data
            
            # Adiciona ao documento
            self.doc.add_picture(img_data, width=Cm(largura_cm))
            return True
            
        except Exception as e:
            logger.warning(f"Erro ao adicionar imagem {url}: {e}")
            return False
    
    def _adicionar_referencias(self, referencias: str):
        """Adiciona referências formatadas."""
        ref_par = self.doc.add_paragraph(referencias)
        ref_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if ref_par.runs:
            ref_run = ref_par.runs[0]
            ref_run.font.size = Pt(self.TAMANHO_FONTE_REFERENCIA)
            ref_run.font.name = "Calibri"
    
    def _adicionar_alternativas(self, alternativas: List[Dict]):
        """Adiciona alternativas da questão."""
        for alt in alternativas:
            self._adicionar_alternativa(alt)
    
    def _adicionar_alternativa(self, alternativa: Dict):
        """Adiciona uma alternativa específica."""
        letra = alternativa.get('letter', '')
        texto = alternativa.get('text', '')
        arquivo = alternativa.get('file', '')
        
        # Alternativa com texto
        if texto:
            adicionar_paragrafo_formatado(self.doc, f"{letra}) {texto}")
        
        # Alternativa com imagem
        if arquivo:
            par = self.doc.add_paragraph()
            par.add_run(f"{letra}) ")
            
            # Tenta adicionar a imagem inline
            try:
                resp = requests.get(arquivo, timeout=10)
                resp.raise_for_status()
                img_data = BytesIO(resp.content)
                
                run = par.add_run()
                run.add_picture(img_data, width=Cm(self.LARGURA_IMAGEM_ALTERNATIVA))
                
            except Exception as e:
                logger.warning(f"Erro ao adicionar imagem da alternativa {letra}: {e}")
    
    def _adicionar_gabarito(self, questoes: List[Questao]):
        """Adiciona seção de gabarito ao documento."""
        # Adiciona quebra de página
        self.doc.add_page_break()
        
        # Título do gabarito
        par = self.doc.add_paragraph()
        run = par.add_run("Gabarito")
        run.bold = True
        run.font.size = Pt(14)
        
        # Lista de respostas
        for questao in questoes:
            gabarito_text = f"{questao.numero}) {questao.gabarito}"
            par = self.doc.add_paragraph(gabarito_text)
            if par.runs:
                par.runs[0].bold = True


class ProcessadorQuestoes:
    """Classe para processar e filtrar questões do banco de dados."""
    
    def __init__(self, pasta_csv: str, pasta_json: str):
        """
        Inicializa o processador.
        
        Args:
            pasta_csv: Caminho para pasta com arquivos CSV
            pasta_json: Caminho para pasta com arquivos JSON
        """
        self.pasta_csv = Path(pasta_csv)
        self.pasta_json = Path(pasta_json)
        self.df_consolidado = None
        self._carregar_csv_consolidado()
    
    def _carregar_csv_consolidado(self):
        """Carrega o arquivo CSV consolidado uma única vez."""
        arquivo = self.pasta_csv / "ITENS_PROVA_CONSOLIDADO.csv"
        
        if not arquivo.exists():
            logger.error(f"Arquivo consolidado nção encontrado: {arquivo}")
            self.df_consolidado = pd.DataFrame()
            return
    
        try:
            self.df_consolidado = pd.read_csv(
                arquivo,
                sep=";",
                encoding="utf-8",
                decimal=","
            )
            
            colunas_numericas = ['CO_HABILIDADE', 'NU_PARAM_A', 'NU_PARAM_B', 'NU_PARAM_C']
            for col in colunas_numericas:
                if col in self.df_consolidado.columns:
                    self.df_consolidado[col] = pd.to_numeric(
                        self.df_consolidado[col],
                        errors='coerce'
                    )
                
            logger.info(f"CSV consolidado carregado: {len(self.df_consolidado)} registros")
        except Exception as e:
            logger.error(f"Erro ao carregar CSV consolidado: {e}")
            self.df_consolidado = pd.DataFrame()
            
            
    def carregar_questoes(self, anos: List[str], materias: List[str], 
                         filtros: Dict) -> pd.DataFrame:
        """
        Carrega e filtra questões de múltiplos anos conforme parâmetros.
        
        Args:
            anos: Lista de anos das provas
            materias: Lista de códigos das matérias
            filtros: Dicionário com filtros adicionais
            
        Returns:
            DataFrame filtrado com as questões de todos os anos
        """
        if self.df_consolidado is None or self.df_consolidado.empty:
            logger.error("DataFrame consolidado não disponível")
            return pd.DataFrame()
        
        # Converte anos para int para comparação
        anos_int = [int(ano) for ano in anos]
        
        # Filtra por anos selecionados
        df_anos = self.df_consolidado[
            self.df_consolidado['ANO_PROVA'].isin(anos_int)
        ].copy()
        
        if df_anos.empty:
            logger.warning(f"Nenhuma questão encontrada para os anos: {anos}")
            return pd.DataFrame()
        
        # Lista para armazenar DataFrames de cada ano
        dataframes_anos = []
        
        # Processa cada ano
        for ano in anos_int:
            df_ano = df_anos[df_anos['ANO_PROVA'] == ano].copy()
            df_ano_filtrado = self._aplicar_filtros(df_ano, str(ano), materias, filtros)
            
            if not df_ano_filtrado.empty:
                df_ano_filtrado['ANO'] = str(ano)
                dataframes_anos.append(df_ano_filtrado)
        
        # Combina todos os DataFrames
        if not dataframes_anos:
            return pd.DataFrame()
    
        df_combinado = pd.concat(dataframes_anos, ignore_index=True)
        
        # Aplica filtros globais
        df_final = self._aplicar_filtros_globais(df_combinado, filtros)
        
        # Adiciona numeraação sequencial
        if not df_final.empty:
             df_final = df_final.reset_index(drop=True)
             df_final["NUMERO_QUESTAO"] = df_final.index + 1
        
        return df_final       
    
    def _aplicar_filtros(self, df: pd.DataFrame, ano: str, 
                        materias: List[str], filtros: Dict) -> pd.DataFrame:
        """Aplica filtros ao DataFrame de um ano específico."""
        # Obtém códigos das provas
        provas_escolhidas = self._obter_provas_escolhidas(ano, materias)
        
        if not provas_escolhidas:
            logger.warning(f"Nenhuma prova encontrada para o ano {ano} e matérias selecionadas")
            return pd.DataFrame()
        
        # Filtro base por matéria e prova
        df_filtrado = df[
            (df["SG_AREA"].isin(materias)) & 
            (df["CO_PROVA"].isin(provas_escolhidas.values()))
        ].copy()
        
        # Filtro por habilidade
        if filtros.get("filtrar_habilidade") and filtros.get("habilidades"):
            df_filtrado = self._filtrar_por_habilidade(df_filtrado, filtros["habilidades"])
        
        # Filtro por posição
        if filtros.get("posicoes"):
            df_filtrado = self._filtrar_por_posicao(df_filtrado, filtros["posicoes"])
        
        # Filtro por dificuldade (se não for filtro por posição)
        elif filtros.get("filtrar_dificuldade", True):
            df_filtrado = self._filtrar_por_dificuldade(
                df_filtrado, 
                filtros.get("min_dif", 0), 
                filtros.get("max_dif", 2000)
            )
        
        return df_filtrado
    
    def _obter_provas_escolhidas(self, ano: str, materias: List[str]) -> Dict:
        """Obtém códigos das provas para as matérias escolhidas."""
        if ano not in provas_regulares:
            logger.warning(f"Ano {ano} não encontrado em provas_regulares")
            return {}
        
        provas = {}
        for materia in materias:
            if materia in provas_regulares[ano]:
                provas[materia] = provas_regulares[ano][materia]
            else:
                logger.warning(f"Matéria {materia} não encontrada para o ano {ano}")
        
        return provas
    
    def _filtrar_por_habilidade(self, 
                                df: pd.DataFrame, 
                                habilidades: List[int],
                                ano: str = None) -> pd.DataFrame:
        """
        Filtra DataFrame por habilidades específicas.
        
        Args:
            df: DataFrame a filtrar
            habilidades: Lista de números de habilidades
            
        Returns:
            DataFrame filtrado
        """
        if df.empty or "CO_HABILIDADE" not in df.columns:
            return df
        
        df_filtrado = df.copy()
        df_filtrado['CO_HABILIDADE'] = pd.to_numeric(
            df_filtrado['CO_HABILIDADE'],
            errors='coerce'
        )
        
        # Converte habilidades para float para comparação
        habilidades_float = [float(h) for h in habilidades]
        
        # Filtra mantendo apenas as habilidades desejadas
        df_filtrado = df_filtrado[df_filtrado["CO_HABILIDADE"].isin(habilidades_float)].copy()
        
        if df_filtrado.empty:
            ano_info = f" do ano {ano}" if ano else ""
            logger.warning(f"Nenhuma questão encontrada para as habilidades: {habilidades}{ano_info}")
        else:
            ano_info = f" no ano {ano}" if ano else ""
            logger.info(f"Encontradas {len(df_filtrado)} questões para as habilidades selecionadas no ano{ano_info}")
        
        return df_filtrado
    
    def _filtrar_por_posicao(self, df: pd.DataFrame, 
                            posicoes_str: str) -> pd.DataFrame:
        """Filtra DataFrame por posições específicas."""
        # Converte string de posições para lista
        posicoes = [int(p.strip()) for p in posicoes_str.split() 
                   if p.strip().isdigit()]
        
        if not posicoes:
            return df
        
        # Filtra e remove duplicatas
        df_filtrado = df[df["CO_POSICAO"].isin(posicoes)].drop_duplicates(
            subset=["CO_POSICAO"]
        ).copy()
        
        # Ordena conforme ordem fornecida
        df_filtrado["CO_POSICAO"] = pd.Categorical(
            df_filtrado["CO_POSICAO"], 
            categories=posicoes, 
            ordered=True
        )
        
        return df_filtrado.sort_values("CO_POSICAO")
    
    def _filtrar_por_dificuldade(self, df: pd.DataFrame, 
                                 min_dif: float, max_dif: float) -> pd.DataFrame:
        """Filtra DataFrame por intervalo de dificuldade."""
        if df.empty:
            return df
        
        # Se não há filtro de dificuldade, apenas ordena
        if min_dif == 0 and max_dif == 2000:
            return df.sort_values("NU_PARAM_B")
        
        # Aplica filtros
        if min_dif == 0:
            df_filtrado = df[df["NU_PARAM_B"] <= max_dif].copy()
        elif max_dif == 2000:
            df_filtrado = df[df["NU_PARAM_B"] >= min_dif].copy()
        else:
            df_filtrado = df[
                (df["NU_PARAM_B"] >= min_dif) & 
                (df["NU_PARAM_B"] <= max_dif)
            ].copy()
        
        return df_filtrado.sort_values("NU_PARAM_B")
    
    def _aplicar_filtros_globais(self, df: pd.DataFrame, filtros: Dict) -> pd.DataFrame:
        """
        Aplica filtros globais ao DataFrame combinado de todos os anos.
        
        Args:
            df: DataFrame combinado
            filtros: Filtros globais
            
        Returns:
            DataFrame com filtros globais aplicados
        """
        if df.empty:
            return df
        
        df_final = df.copy()
        
        # Aplica filtro de quantidade se especificado
        if filtros.get("filtrar_quantidade") and filtros.get("quantidade"):
            quantidade_max = filtros["quantidade"]
            if len(df_final) > quantidade_max:
                # Se já está ordenado por dificuldade, mantém ordem
                # Caso contrário, ordena por ano e dificuldade
                if not filtros.get("posicoes"):
                    df_final = df_final.sort_values(["ANO", "NU_PARAM_B"])
                
                df_final = df_final.head(quantidade_max)
                logger.info(f"Limitado a {quantidade_max} questões conforme solicitado")
        
        return df_final
    
    def carregar_detalhes_questao(self, ano: int, posicao: int) -> Optional[Dict]:
        """
        Carrega detalhes de uma questão do arquivo JSON.
        
        Args:
            ano: Ano da questão
            posicao: Posição/número da questão
            
        Returns:
            Dicionário com detalhes ou None se não encontrado
        """
        arquivo = self.pasta_json / str(ano) / "questions" / str(posicao) / "details.json"
        
        if not arquivo.exists():
            logger.warning(f"JSON não encontrado: {arquivo}")
            return None
        
        try:
            with open(arquivo, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Erro ao carregar JSON: {e}")
            return None
    
    def criar_objeto_questao(self, row: pd.Series, dados_json: Dict, 
                           numero: int, ano: int) -> Questao:
        """
        Cria objeto Questao a partir dos dados.
        
        Args:
            row: Linha do DataFrame
            dados_json: Dados do JSON
            numero: Número sequencial da questão
            ano: Ano da prova
            
        Returns:
            Objeto Questao preenchido
        """
        # Processa habilidade
        habilidade = "-"
        if not pd.isna(row.get("CO_HABILIDADE")):
            try:
                habilidade = str(int(float(row["CO_HABILIDADE"])))
            except:
                pass
        
        return Questao(
            numero=numero,
            ano=ano,
            posicao=int(row["CO_POSICAO"]),
            nota=float(row["NU_PARAM_B"]),
            habilidade=habilidade,
            gabarito=row["TX_GABARITO"],
            contexto=dados_json.get("context"),
            referencias=dados_json.get("references"),
            introducao_alternativas=dados_json.get("alternativesIntroduction"),
            alternativas=dados_json.get("alternatives", []),
            arquivos=dados_json.get("files", [])
        )


class AplicacaoENEM:
    """Classe principal da aplicação."""
    
    def __init__(self):
        """Inicializa a aplicação."""
        self.processador = ProcessadorQuestoes(
            pasta_csv="base-de-dados-CSV",
            pasta_json=os.path.join("enem-api", "public")
        )
        self.gerador = GeradorDocumento()
    
    def executar(self):
        """Executa a aplicação principal."""
        st.title("🎓 Gerador de Compilado de Questões ENEM")
        
        # Obtém entradas do usuário
        (anos, materias, posicoes, max_dif, min_dif, 
         filtrar_quantidade, quantidade, filtrar_habilidade, habilidades) = entradas()
        
        # Botão de geração
        if st.button("📄 Gerar DOCX", type="primary"):
            self._gerar_documento(
                anos, materias, posicoes, min_dif, max_dif, 
                filtrar_quantidade, quantidade, filtrar_habilidade, habilidades
            )
    
    def _gerar_documento(self, anos: List[str], materias: List[str], 
                         posicoes: Optional[str], min_dif: int, max_dif: int,
                         filtrar_quantidade: bool = False, 
                         quantidade: Optional[int] = None,
                         filtrar_habilidade: bool = False,
                         habilidades: Optional[List[int]] = None):
        """Gera o documento com as questões filtradas."""
        
        # Validações iniciais
        if not anos:
            st.error("❌ Selecione pelo menos um ano!")
            return
        
        if not materias:
            st.error("❌ Selecione pelo menos uma matéria!")
            return
        
        # Prepara filtros
        filtros = {
            "posicoes": posicoes,
            "min_dif": min_dif,
            "max_dif": max_dif,
            "filtrar_quantidade": filtrar_quantidade,
            "quantidade": quantidade,
            "filtrar_dificuldade": not bool(posicoes),  # Só filtra por dificuldade se não for por posição
            "filtrar_habilidade": filtrar_habilidade,
            "habilidades": habilidades
        }
        
        # Carrega questões
        with st.spinner("📊 Carregando questões..."):
            df_questoes = self.processador.carregar_questoes(anos, materias, filtros)
        
        if df_questoes.empty:
            st.warning("⚠️ Nenhuma questão encontrada com os filtros selecionados.")
            return
        
        # Mostra preview
        st.info(f"📋 {len(df_questoes)} questão(ões) encontrada(s)")
        
        # Processa questões
        questoes = self._processar_questoes(df_questoes)
        
        if not questoes:
            st.error("❌ Erro ao processar questões.")
            return
        
        # Gera documento
        with st.spinner("📝 Gerando documento..."):
            doc = self.gerador.criar_documento(questoes, anos)
        
        # Salva e oferece download
        self._salvar_e_baixar(doc, anos, len(questoes))
    
    def _processar_questoes(self, df: pd.DataFrame) -> List[Questao]:
        """Processa DataFrame e cria lista de objetos Questao."""
        questoes = []
        total = len(df)
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        for idx, (_, row) in enumerate(df.iterrows(), start=1):
            # Atualiza progresso
            progress = idx / total
            progress_bar.progress(progress)
            status_text.text(f"Processando questão {idx} de {total}...")
            
            # Obtém ano da linha
            ano = int(row["ANO"]) if "ANO" in row else int(row.name)
            
            # Carrega dados JSON
            dados_json = self.processador.carregar_detalhes_questao(
                ano, 
                int(row["CO_POSICAO"])
            )
            
            if dados_json:
                questao = self.processador.criar_objeto_questao(
                    row, dados_json, idx, ano
                )
                questoes.append(questao)
            else:
                logger.warning(f"Dados JSON não encontrados para questão {ano}-{int(row['CO_POSICAO'])}")
        
        progress_bar.empty()
        status_text.empty()
        
        return questoes
    
    def _salvar_e_baixar(self, doc: Document, anos: List[str], num_questoes: int):
        """Salva documento e oferece para download."""
        # Cria nome do arquivo
        if len(anos) == 1:
            nome_arquivo = f"ENEM_{anos[0]}_{num_questoes}_questoes.docx"
        else:
            anos_ordenados = sorted(anos)
            anos_texto = f"{anos_ordenados[0]}-{anos_ordenados[-1]}"
            nome_arquivo = f"ENEM_{anos_texto}_{num_questoes}_questoes.docx"
        
        # Salva em buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Oferece download
        st.success(f"✅ Documento gerado com {num_questoes} questões!")
        
        st.download_button(
            label="📥 Baixar DOCX",
            data=buffer.getvalue(),
            file_name=nome_arquivo,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )


def main():
    """Função principal."""
    # Configuração da página
    st.set_page_config(
        page_title="Gerador ENEM",
        page_icon="📚",
        layout="wide"
    )
    
    # Executa aplicação
    app = AplicacaoENEM()
    app.executar()


if __name__ == "__main__":
    main()