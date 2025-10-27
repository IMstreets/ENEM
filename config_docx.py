# Configurações do documento
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
import re
from typing import Optional, Dict, Any


class DocumentFormatter:
    """Classe para gerenciar formatação de documentos Word."""
    
    # Configurações padrão
    DEFAULT_FONT = "Calibri"
    DEFAULT_FONT_SIZE = 11
    DEFAULT_MARGINS = {"top": 1, "bottom": 2.5, "left": 1, "right": 1}
    DEFAULT_COLUMNS = {"num": 2, "space": 200, "sep": True}
    DEFAULT_BORDER = {"style": "single", "size": 2, "space": 5, "color": "000000"}
    
    def __init__(self, doc):
        """Inicializa o formatador com um documento."""
        self.doc = doc
        self.section = doc.sections[0]
        self.sectPr = self.section._sectPr
    
    def configurar_documento(self) -> Any:
        """Aplica todas as configurações ao documento."""
        self._configurar_fonte()
        self._configurar_alinhamento()
        self._configurar_margens()
        self._configurar_colunas()
        self._configurar_bordas()
        return self.doc
    
    def _configurar_fonte(self):
        """Define a fonte padrão do documento."""
        style = self.doc.styles["Normal"]
        style.font.name = self.DEFAULT_FONT
        style.font.size = Pt(self.DEFAULT_FONT_SIZE)
    
    def _configurar_alinhamento(self):
        """Configura alinhamento justificado para estilos principais."""
        estilos_para_justificar = ["Normal", "Heading 1", "Heading 2"]
        for style_name in estilos_para_justificar:
            if style_name in self.doc.styles:
                style = self.doc.styles[style_name]
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _configurar_margens(self):
        """Define as margens do documento."""
        self.section.top_margin = Cm(self.DEFAULT_MARGINS["top"])
        self.section.bottom_margin = Cm(self.DEFAULT_MARGINS["bottom"])
        self.section.left_margin = Cm(self.DEFAULT_MARGINS["left"])
        self.section.right_margin = Cm(self.DEFAULT_MARGINS["right"])
    
    def _configurar_colunas(self):
        """Configura as colunas do documento."""
        # Remove configurações anteriores de colunas
        for child in self.sectPr.findall(ns.qn("w:cols")):
            self.sectPr.remove(child)
        
        # Adiciona nova configuração de colunas
        cols = self._criar_elemento_colunas()
        self.sectPr.append(cols)
    
    def _criar_elemento_colunas(self) -> OxmlElement:
        """Cria elemento XML para configuração de colunas."""
        cols = OxmlElement("w:cols")
        cols.set(ns.qn("w:num"), str(self.DEFAULT_COLUMNS["num"]))
        cols.set(ns.qn("w:space"), str(self.DEFAULT_COLUMNS["space"]))
        cols.set(ns.qn("w:sep"), str(self.DEFAULT_COLUMNS["sep"]).lower())
        return cols
    
    def _configurar_bordas(self):
        """Adiciona bordas à página."""
        pgBorders = OxmlElement("w:pgBorders")
        
        for lado in ["top", "left", "bottom", "right"]:
            borda = self._criar_elemento_borda(lado)
            pgBorders.append(borda)
        
        self.sectPr.append(pgBorders)
    
    def _criar_elemento_borda(self, lado: str) -> OxmlElement:
        """Cria elemento XML para uma borda específica."""
        borda = OxmlElement(f"w:{lado}")
        borda.set(ns.qn("w:val"), self.DEFAULT_BORDER["style"])
        borda.set(ns.qn("w:sz"), str(self.DEFAULT_BORDER["size"]))
        borda.set(ns.qn("w:space"), str(self.DEFAULT_BORDER["space"]))
        borda.set(ns.qn("w:color"), self.DEFAULT_BORDER["color"])
        return borda


class ParagraphFormatter:
    """Classe para formatação de parágrafos com estilos especiais."""
    
    # Padrões de regex para formatação
    PATTERNS = {
        "bold": (r"\*\*(?:.*?)\*\*", lambda run: setattr(run, 'bold', True)),
        "subscript": (r"__(?:.*?)__", lambda run: setattr(run.font, 'subscript', True)),
        "italic": (r"_(?:.*?)_", lambda run: setattr(run, 'italic', True)),
        "superscript": (r"\^(?:.*?)\^", lambda run: setattr(run.font, 'superscript', True))
    }
    
    def __init__(self, doc):
        """Inicializa o formatador de parágrafos."""
        self.doc = doc
        self._compile_patterns()
    
    def _compile_patterns(self):
        """Compila todos os padrões de regex em um único padrão."""
        patterns = [f"({pattern})" for pattern, _ in self.PATTERNS.values()]
        self.compiled_pattern = re.compile("|".join(patterns))
    
    def adicionar_paragrafo_formatado(self, texto: str) -> Any:
        """
        Adiciona um parágrafo formatado ao documento.
        
        Suporta:
        - **negrito**
        - _itálico_
        - ^sobrescrito^
        - __subscrito__
        - Listas com marcadores (* -)
        """
        texto_strip = texto.strip()
        
        # Determina o estilo do parágrafo
        par = self._criar_paragrafo_base(texto_strip)
        
        # Processa o conteúdo com formatação
        conteudo = self._extrair_conteudo(texto_strip)
        self._aplicar_formatacao(par, conteudo)
        
        return par
    
    def _criar_paragrafo_base(self, texto: str) -> Any:
        """Cria o parágrafo base com o estilo apropriado."""
        if self._eh_lista_marcador(texto):
            return self.doc.add_paragraph(style="List Bullet")
        return self.doc.add_paragraph()
    
    def _eh_lista_marcador(self, texto: str) -> bool:
        """Verifica se o texto é uma lista com marcador."""
        return texto.startswith("* \\-") or texto.startswith("* -")
    
    def _extrair_conteudo(self, texto: str) -> str:
        """Extrai o conteúdo limpo do texto."""
        if self._eh_lista_marcador(texto):
            return re.sub(r"^\*\s*\\?-", "", texto).strip()
        return texto
    
    def _aplicar_formatacao(self, par: Any, texto: str):
        """Aplica formatação ao parágrafo baseado nos padrões."""
        last_index = 0
        
        for match in self.compiled_pattern.finditer(texto):
            start, end = match.span()
            
            # Adiciona texto normal antes do match
            if start > last_index:
                par.add_run(texto[last_index:start])
            
            # Aplica formatação específica
            self._processar_match(par, match)
            last_index = end
        
        # Adiciona texto restante
        if last_index < len(texto):
            par.add_run(texto[last_index:])
    
    def _processar_match(self, par: Any, match: re.Match):
        """Processa um match de formatação e aplica o estilo apropriado."""
        # Mapeia grupos do regex para formatações
        formatacoes = {
            (1, lambda r: setattr(r, 'bold', True)),           # **negrito**
            (2, lambda r: setattr(r.font, 'subscript', True)), # __subscrito__
            (3, lambda r: setattr(r, 'italic', True)),         # _itálico_
            (4, lambda r: setattr(r.font, 'superscript', True)) # ^sobrescrito^
        }
        
        for grupo, aplicar_formato in formatacoes:
            if match.group(grupo):
                texto = re.sub(r"(^\*\*|\*\*$|^__|__$|^_|_$|^\^|\^$)", "", match.group(grupo))
                run = par.add_run(texto)
                aplicar_formato(run)
                break


# Funções de conveniência para manter compatibilidade
def configurar_documento(doc) -> Any:
    """Função de compatibilidade para configurar documento."""
    formatter = DocumentFormatter(doc)
    return formatter.configurar_documento()


def adicionar_paragrafo_formatado(doc, texto: str) -> Any:
    """Função de compatibilidade para adicionar parágrafo formatado."""
    formatter = ParagraphFormatter(doc)
    return formatter.adicionar_paragrafo_formatado(texto)